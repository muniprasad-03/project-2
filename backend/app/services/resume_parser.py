from __future__ import annotations

import logging
from io import BytesIO
from typing import Union

import pdfplumber
from docx import Document


logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Supported file types
# ---------------------------------------------------------------------------

SUPPORTED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
}


# ---------------------------------------------------------------------------
# Exceptions
# ---------------------------------------------------------------------------

class ResumeParserError(Exception):
    """Base exception for resume parsing errors."""


class UnsupportedResumeTypeError(ResumeParserError):
    """Raised when the uploaded resume type is not supported."""


class EmptyResumeError(ResumeParserError):
    """Raised when no usable text can be extracted."""


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

MAX_RESUME_SIZE_BYTES = 5 * 1024 * 1024


# ---------------------------------------------------------------------------
# Validation
# ---------------------------------------------------------------------------

def validate_resume_file(
    file_bytes: bytes,
    filename: str,
    content_type: str | None = None,
) -> None:
    """
    Validate resume size and file type.

    Supported:
        PDF
        DOCX

    Maximum size:
        5 MB
    """

    if not file_bytes:
        raise ResumeParserError(
            "The uploaded resume is empty."
        )

    if len(file_bytes) > MAX_RESUME_SIZE_BYTES:
        raise ResumeParserError(
            "Resume file exceeds the maximum allowed size of 5 MB."
        )

    filename_lower = filename.lower()

    extension_supported = any(
        filename_lower.endswith(extension)
        for extension in SUPPORTED_EXTENSIONS
    )

    content_type_supported = (
        content_type in SUPPORTED_CONTENT_TYPES
        if content_type
        else False
    )

    if not extension_supported and not content_type_supported:
        raise UnsupportedResumeTypeError(
            "Unsupported resume format. "
            "Only PDF and DOCX files are supported."
        )


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------

def _extract_pdf_text(
    file_bytes: bytes,
) -> str:
    """
    Extract text from a PDF entirely in memory.
    """

    text_parts = []

    try:
        with pdfplumber.open(
            BytesIO(file_bytes)
        ) as pdf:

            for page_number, page in enumerate(
                pdf.pages,
                start=1,
            ):
                try:
                    page_text = page.extract_text()

                except Exception as exc:
                    logger.warning(
                        "Unable to extract text from PDF page %d: %s",
                        page_number,
                        exc,
                    )
                    continue

                if page_text:
                    text_parts.append(
                        page_text
                    )

    except Exception as exc:
        raise ResumeParserError(
            "Unable to read the PDF resume."
        ) from exc

    return "\n".join(text_parts)


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def _extract_docx_text(
    file_bytes: bytes,
) -> str:
    """
    Extract text from a DOCX file entirely in memory.
    """

    try:
        document = Document(
            BytesIO(file_bytes)
        )

    except Exception as exc:
        raise ResumeParserError(
            "Unable to read the DOCX resume."
        ) from exc

    text_parts = []

    # ---------------------------------------------------------------
    # Paragraphs
    # ---------------------------------------------------------------

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:
            text_parts.append(text)

    # ---------------------------------------------------------------
    # Tables
    # ---------------------------------------------------------------

    for table in document.tables:

        for row in table.rows:

            row_parts = []

            for cell in row.cells:

                cell_text = cell.text.strip()

                if cell_text:
                    row_parts.append(cell_text)

            if row_parts:
                text_parts.append(
                    " | ".join(row_parts)
                )

    return "\n".join(text_parts)


# ---------------------------------------------------------------------------
# Text cleanup
# ---------------------------------------------------------------------------

def _clean_text(
    text: str,
) -> str:
    """
    Normalize extracted resume text.
    """

    if not text:
        return ""

    lines = []

    for line in text.splitlines():

        line = " ".join(
            line.split()
        )

        if line:
            lines.append(line)

    return "\n".join(lines).strip()


# ---------------------------------------------------------------------------
# Main extraction function
# ---------------------------------------------------------------------------

def extract_text(
    file_bytes: bytes,
    content_type: str | None = None,
    filename: str = "resume",
) -> str:
    """
    Extract text from a PDF or DOCX resume.

    The file is processed entirely in memory.

    Parameters
    ----------
    file_bytes:
        Raw uploaded file bytes.

    content_type:
        MIME type supplied by the client.

    filename:
        Original filename.

    Returns
    -------
    str
        Extracted and cleaned resume text.
    """

    validate_resume_file(
        file_bytes=file_bytes,
        filename=filename,
        content_type=content_type,
    )

    filename_lower = filename.lower()

    # ---------------------------------------------------------------
    # Determine file type
    # ---------------------------------------------------------------

    is_pdf = (
        filename_lower.endswith(".pdf")
        or content_type == "application/pdf"
    )

    is_docx = (
        filename_lower.endswith(".docx")
        or content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # ---------------------------------------------------------------
    # Extract
    # ---------------------------------------------------------------

    if is_pdf:

        text = _extract_pdf_text(
            file_bytes
        )

    elif is_docx:

        text = _extract_docx_text(
            file_bytes
        )

    else:
        raise UnsupportedResumeTypeError(
            "Unable to determine the resume file type."
        )

    # ---------------------------------------------------------------
    # Clean
    # ---------------------------------------------------------------

    text = _clean_text(text)

    if not text:
        raise EmptyResumeError(
            "No readable text could be extracted from the resume."
        )

    logger.info(
        "Resume parsed successfully: %s (%d characters)",
        filename,
        len(text),
    )

    return text


# ---------------------------------------------------------------------------
# Convenience helpers
# ---------------------------------------------------------------------------

def extract_text_from_pdf(
    file_bytes: bytes,
) -> str:
    """
    Convenience wrapper for PDF files.
    """

    return extract_text(
        file_bytes=file_bytes,
        content_type="application/pdf",
        filename="resume.pdf",
    )


def extract_text_from_docx(
    file_bytes: bytes,
) -> str:
    """
    Convenience wrapper for DOCX files.
    """

    return extract_text(
        file_bytes=file_bytes,
        content_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        filename="resume.docx",
    )